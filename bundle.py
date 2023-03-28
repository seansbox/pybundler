import os
import sys
from glob import glob
from docx import Document
from docx.shared import Inches, Pt
from termcolor import colored

# https://community.canvaslms.com/t5/Instructor-Guide/What-types-of-files-can-be-previewed-in-Canvas/ta-p/607

if getattr(sys, 'frozen', False):
    app_path = os.path.dirname(sys.executable)
elif __file__:
    app_path = os.path.dirname(__file__)
os.chdir(app_path)

types = [
    '**/*.md', '**/Pipfile', '**/Procfile',
    '**/*.py', '**/*.htm', '**/*.html', '**/*.css',
    '**/*.csv', '**/*.json', '**/*.xml',
    '**/*.png', '**/*.jpg', '**/*.gif',
]

files = []
for t in types: files.extend(glob(t, recursive=True))
files = sorted(files)

document = Document()
sections = document.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

print("Bundling files...\n")
for file in files:
    if 'build' in file or 'dist' in file or '__' in file: continue
    if 'manage.py' in file or 'asgi.py' in file or 'wsgi.py' in file: continue
    print(f" - {file}")
    h = document.add_heading(file, 0)
    h.style.font.size = Pt(16)
    h.style.font.bold = True
    if '.png' in file.lower() or '.jpg' in file.lower() or '.gif' in file.lower():
        document.add_picture(file, width=Inches(7.5))
    else:
        with open(file, 'r', encoding='utf-8', errors='ignore') as f:
            p = document.add_paragraph(f.read())
            p.style.font.size = Pt(8)
            p.style.font.name = 'Courier New'
            p.style.font.bold = True
    document.add_page_break()

document.save('bundle.docx')

input("\nPress Enter to close... ")
